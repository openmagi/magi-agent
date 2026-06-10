"""DocumentWrite DOCX backend — render markdown source into a ``.docx`` file.

This module is intentionally lazy-import only with respect to ``python-docx``:
``docx`` is imported *inside* :func:`docx_write` so that importing this module
(or :mod:`magi_agent.plugins.native.documents`) never pulls ``docx`` into
``sys.modules``.  When ``python-docx`` is not installed the handler returns
``status="blocked"`` with ``errorCode="document_dependency_not_installed"``.

``python-docx>=1.1`` ships in the optional ``files`` extra
(``uv sync --extra files``); DOCX availability is therefore naturally gated by
that optional dependency rather than an env flag.

The markdown parser is deliberately pragmatic.  Its contract is that the literal
source text lands in the document (Task B verifies coverage).  Supported
constructs:

* ATX headings ``#``..``######`` -> ``add_heading(text, level)``
* Blank-line-separated paragraphs -> ``add_paragraph``
* Unordered list items (``- ``/``* ``) -> ``add_paragraph(..., style="List Bullet")``
* Ordered list items (``1. ``) -> ``add_paragraph(..., style="List Number")``
* Fenced code blocks (```` ``` ````) -> a monospace paragraph (content preserved)
* Simple pipe tables (``| a | b |`` + ``---`` separator) -> ``add_table``
* Minimal inline emphasis (``**bold**`` / ``*italic*``)
* Unknown lines -> plain paragraphs (no words dropped)

Redaction / coverage contract
------------------------------
The document is rendered from ``redact_public_text(source)`` — the same
redaction applied on the markdown path in
:mod:`magi_agent.plugins.native.documents`.  File paths (``/home/…``,
``/Users/…``, etc.), private-line patterns, and other sensitive tokens are
replaced or dropped *before* any content reaches the ``python-docx`` renderer.

**Task B (coverage verifier) MUST compare against the redacted source**, not the
raw input.  Comparing against the raw ``source`` will produce false coverage
failures because the redacted tokens (e.g. ``[redacted-path]``) will not match
the original path strings.  The correct reference text is::

    from magi_agent.web_acquisition.policy import redact_public_text
    reference = redact_public_text(source, max_chars=200_000)

Do NOT change the redaction behaviour here — keep it consistent with the
markdown path so both output formats share the same redaction contract.
"""

from __future__ import annotations

import hashlib
import io
import re

from magi_agent.plugins.native._common import (
    blocked_result,
    digest,
    safe_child_path,
)
from magi_agent.tools.context import ToolContext
from magi_agent.tools.document_write.model import MIME_TYPES, PREVIEW_KINDS
from magi_agent.tools.result import ToolResult
from magi_agent.web_acquisition.policy import redact_public_text

_DEFAULT_NAME = "magi-document.docx"
_MAX_CHARS = 200_000

_HEADING_RE = re.compile(r"^(#{1,6})\s+(.*)$")
_UNORDERED_RE = re.compile(r"^[-*]\s+(.*)$")
_ORDERED_RE = re.compile(r"^\d+[.)]\s+(.*)$")
_FENCE_RE = re.compile(r"^\s*```")
_TABLE_SEP_RE = re.compile(r"^\s*\|?\s*:?-{1,}:?\s*(\|\s*:?-{1,}:?\s*)+\|?\s*$")


def docx_write(arguments: dict[str, object], context: ToolContext) -> ToolResult:
    """Render markdown ``content`` into a ``.docx`` file inside the workspace."""
    source = str(
        arguments.get("content")
        or arguments.get("text")
        or arguments.get("source")
        or arguments.get("markdown")
        or ""
    )
    if not source.strip():
        return blocked_result("DocumentWrite", "content_required")

    # Lazy import — never a top-level ``import docx``.
    try:
        from docx import Document  # noqa: PLC0415
    except ImportError:
        return blocked_result("DocumentWrite", "document_dependency_not_installed")

    path_value = arguments.get("path") or arguments.get("filename") or _DEFAULT_NAME
    try:
        path = safe_child_path(
            context,
            path_value,
            default_name=_DEFAULT_NAME,
            mutating=True,
        )
    except ValueError as error:
        return blocked_result("DocumentWrite", str(error))

    safe_source = redact_public_text(source, max_chars=_MAX_CHARS)

    try:
        document = Document()
        _render_markdown(document, safe_source)
        buf = io.BytesIO()
        document.save(buf)
        data = buf.getvalue()
        path.parent.mkdir(parents=True, exist_ok=True)
        path.write_bytes(data)
    except Exception:  # noqa: BLE001 — never raise out of the handler.
        return blocked_result("DocumentWrite", "document_write_failed")

    relative = path.relative_to(
        safe_child_path(context, ".", default_name=".", mutating=False)
    ).as_posix()
    content_digest = "sha256:" + hashlib.sha256(data).hexdigest()
    short_digest = content_digest.removeprefix("sha256:")[:16]
    artifact_ref = "artifact:docx:" + short_digest

    # Deterministic source-content coverage verifier (audit-only here; Task C
    # adds blocking). Compare against the REDACTED ``safe_source`` — never the raw
    # input — per the module-level redaction/coverage contract. Fail-open: any
    # error in coverage building must never fail a successful write.
    coverage_projection, evidence_declaration = _build_coverage_evidence(
        document=document,
        safe_source=safe_source,
    )

    output: dict[str, object] = {
        "path": relative,
        "pathRef": relative,
        "contentDigest": content_digest,
        "byteCount": len(data),
        "format": "docx",
        "mimeType": MIME_TYPES["docx"],
        "previewKind": PREVIEW_KINDS["docx"],
        "localOnly": True,
        "artifactRef": artifact_ref,
        "artifactRefs": (artifact_ref,),
    }
    if coverage_projection is not None:
        output["coverage"] = coverage_projection

    metadata: dict[str, object] = {
        "toolName": "DocumentWrite",
        "handler": "first_party_native_local",
        "outputDigest": digest(output),
    }
    if evidence_declaration is not None:
        # Established emission channel: ``LocalToolEvidenceCollector`` →
        # ``evidence_from_tool_result`` reads ``metadata["evidence"]`` and builds
        # the canonical DocumentCoverage ``EvidenceRecord`` consumed by the
        # verifier-bus (Task C).
        metadata["evidence"] = evidence_declaration

    return ToolResult(
        status="ok",
        output=output,
        llmOutput=output,
        transcriptOutput={
            "toolName": "DocumentWrite",
            "outputDigest": digest(output),
        },
        artifactRefs=(artifact_ref,),
        metadata=metadata,
    )


def _build_coverage_evidence(
    *,
    document: object,
    safe_source: str,
) -> tuple[dict[str, object] | None, dict[str, object] | None]:
    """Build the coverage projection + evidence declaration. Fail-open.

    Returns ``(None, None)`` if anything goes wrong so a successful write is never
    failed by the audit-only coverage step.
    """
    try:
        from magi_agent.evidence.document_coverage import (  # noqa: PLC0415
            DocumentCoverageBoundary,
            evidence_declaration_from_record,
        )
        from magi_agent.tools.document_tools import extract_docx_text  # noqa: PLC0415

        rendered_text = extract_docx_text(document)
        record = DocumentCoverageBoundary().build_record(
            source_markdown=safe_source,
            doc_text=rendered_text,
        )
        declaration = evidence_declaration_from_record(
            record,
            tool_name="DocumentWrite",
        )
        return record.public_projection(), declaration
    except Exception:  # noqa: BLE001 — coverage is audit-only; never fail the write.
        return None, None


# ---------------------------------------------------------------------------
# Markdown -> docx rendering
# ---------------------------------------------------------------------------


def _render_markdown(document: object, source: str) -> None:
    lines = source.replace("\r\n", "\n").replace("\r", "\n").split("\n")
    index = 0
    total = len(lines)
    while index < total:
        line = lines[index]
        stripped = line.strip()

        # Fenced code block.
        if _FENCE_RE.match(line):
            index += 1
            code_lines: list[str] = []
            while index < total and not _FENCE_RE.match(lines[index]):
                code_lines.append(lines[index])
                index += 1
            if index < total:  # consume closing fence
                index += 1
            _add_code_block(document, "\n".join(code_lines))
            continue

        # Blank line — skip.
        if not stripped:
            index += 1
            continue

        # Heading.
        heading_match = _HEADING_RE.match(stripped)
        if heading_match:
            level = len(heading_match.group(1))
            document.add_heading(heading_match.group(2).strip(), level=level)
            index += 1
            continue

        # Pipe table: current line looks like a row and the next is a separator.
        if (
            "|" in line
            and index + 1 < total
            and _TABLE_SEP_RE.match(lines[index + 1])
        ):
            consumed = _add_table(document, lines, index)
            index += consumed
            continue

        # Unordered list item.
        unordered = _UNORDERED_RE.match(stripped)
        if unordered:
            _add_styled_paragraph(document, unordered.group(1), "List Bullet")
            index += 1
            continue

        # Ordered list item.
        ordered = _ORDERED_RE.match(stripped)
        if ordered:
            _add_styled_paragraph(document, ordered.group(1), "List Number")
            index += 1
            continue

        # Plain paragraph — gather consecutive non-blank, non-structural lines.
        para_lines: list[str] = []
        while index < total:
            candidate = lines[index]
            candidate_stripped = candidate.strip()
            if not candidate_stripped:
                break
            if _FENCE_RE.match(candidate) or _HEADING_RE.match(candidate_stripped):
                break
            if _UNORDERED_RE.match(candidate_stripped) or _ORDERED_RE.match(
                candidate_stripped
            ):
                break
            if (
                "|" in candidate
                and index + 1 < total
                and _TABLE_SEP_RE.match(lines[index + 1])
            ):
                break
            para_lines.append(candidate_stripped)
            index += 1
        _add_paragraph_with_inline(document, " ".join(para_lines))


def _add_styled_paragraph(document: object, text: str, style: str) -> None:
    try:
        paragraph = document.add_paragraph(style=style)
    except (KeyError, ValueError):
        # Style missing from the default template — fall back to plain text.
        paragraph = document.add_paragraph()
    _apply_inline(paragraph, text)


def _add_paragraph_with_inline(document: object, text: str) -> None:
    paragraph = document.add_paragraph()
    _apply_inline(paragraph, text)


def _add_code_block(document: object, code: str) -> None:
    paragraph = document.add_paragraph()
    run = paragraph.add_run(code)
    try:
        run.font.name = "Courier New"
    except Exception:  # noqa: BLE001 — cosmetic only.
        pass


def _split_table_row(line: str) -> list[str]:
    trimmed = line.strip()
    if trimmed.startswith("|"):
        trimmed = trimmed[1:]
    if trimmed.endswith("|"):
        trimmed = trimmed[:-1]
    return [cell.strip() for cell in trimmed.split("|")]


def _add_table(document: object, lines: list[str], start: int) -> int:
    """Render a pipe table starting at ``start``; return the line count consumed."""
    header = _split_table_row(lines[start])
    # lines[start + 1] is the separator (already matched by caller).
    body: list[list[str]] = []
    index = start + 2
    total = len(lines)
    while index < total and "|" in lines[index] and lines[index].strip():
        body.append(_split_table_row(lines[index]))
        index += 1

    width = max([len(header)] + [len(row) for row in body], default=len(header))
    table = document.add_table(rows=0, cols=width)
    try:
        table.style = "Table Grid"
    except (KeyError, ValueError):
        pass

    for source_row in [header, *body]:
        cells = table.add_row().cells
        for col in range(width):
            value = source_row[col] if col < len(source_row) else ""
            _apply_inline(cells[col].paragraphs[0], value)
    return index - start


# ---------------------------------------------------------------------------
# Minimal inline emphasis
# ---------------------------------------------------------------------------

_INLINE_RE = re.compile(r"(\*\*[^*]+\*\*|\*[^*]+\*)")


def _apply_inline(paragraph: object, text: str) -> None:
    if not text:
        paragraph.add_run("")
        return
    for token in _INLINE_RE.split(text):
        if not token:
            continue
        if token.startswith("**") and token.endswith("**") and len(token) > 4:
            run = paragraph.add_run(token[2:-2])
            run.bold = True
        elif token.startswith("*") and token.endswith("*") and len(token) > 2:
            run = paragraph.add_run(token[1:-1])
            run.italic = True
        else:
            paragraph.add_run(token)


__all__ = ["docx_write"]

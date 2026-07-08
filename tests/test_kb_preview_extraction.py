"""KB file preview extracts binary documents instead of returning not_found.

The dashboard KB panel previews a file via ``GET /v1/app/knowledge/file``. It
used to ``read_text(utf-8)`` the file and return ``{"error": "not_found"}`` on
``UnicodeDecodeError`` — so every binary document (PDF, Office, legacy Excel)
previewed as a misleading red ``not_found`` even though the file was present and
listed. These tests pin the fix: text-like files still fast-read, binary
documents are extracted to markdown via the first-party converter, and a genuine
extraction failure yields an honest, specific error (never ``not_found``).
"""

from __future__ import annotations

import tomllib
from pathlib import Path

from fastapi.testclient import TestClient

from magi_agent.app import create_app
from magi_agent.config.models import BuildInfo, RuntimeConfig
from magi_agent.runtime.openmagi_runtime import OpenMagiRuntime

_TOKEN = "local-dev-token"
_WORKSPACE_ENV_VARS = (
    "MAGI_WORKSPACE_ROOT",
    "MAGI_WORKSPACE",
    "MAGI_AGENT_WORKSPACE",
    "CORE_AGENT_PYTHON_GATE5B_FULL_TOOLHOST_WORKSPACE_ROOT",
    "CORE_AGENT_WORKSPACE_ROOT",
)


def _runtime() -> OpenMagiRuntime:
    return OpenMagiRuntime(
        config=RuntimeConfig(
            bot_id="local-bot",
            user_id="local-user",
            gateway_token=_TOKEN,
            api_proxy_url="http://api-proxy.local",
            chat_proxy_url="http://chat-proxy.local",
            redis_url="redis://redis.local:6379/0",
            model="gpt-5.2",
            build=BuildInfo(version="0.1.0", build_sha="sha-test"),
        )
    )


def _client(tmp_path, monkeypatch) -> TestClient:
    monkeypatch.chdir(tmp_path)
    monkeypatch.setenv("MAGI_CONFIG", str(tmp_path / "config.toml"))
    for name in _WORKSPACE_ENV_VARS:
        monkeypatch.delenv(name, raising=False)
    client = TestClient(create_app(_runtime()))
    client.headers.update({"x-gateway-token": _TOKEN})
    return client


def _kb_dir(tmp_path) -> Path:
    # The upload endpoint writes to knowledge/<collection>/<filename>; the
    # listing returns workspace-relative POSIX paths the preview then re-resolves.
    downloads = tmp_path / "knowledge" / "Downloads"
    downloads.mkdir(parents=True, exist_ok=True)
    return downloads


def _preview(client: TestClient, rel: str):
    from urllib.parse import quote

    return client.get(f"/v1/app/knowledge/file?path={quote(rel)}")


# --------------------------------------------------------------------------- #
# Binary documents are extracted (the reported bug)
# --------------------------------------------------------------------------- #
def test_docx_preview_extracts_text(tmp_path, monkeypatch) -> None:
    from docx import Document

    target = _kb_dir(tmp_path) / "report.docx"
    doc = Document()
    doc.add_paragraph("NAEOE_DISTILLERY financial summary line")
    doc.save(str(target))

    client = _client(tmp_path, monkeypatch)
    res = _preview(client, "knowledge/Downloads/report.docx")

    assert res.status_code == 200, res.text
    body = res.json()
    assert "NAEOE_DISTILLERY financial summary line" in body["content"]
    assert body["extracted"] is True
    assert body["sourceTool"] == "document_read"


def test_xlsx_preview_extracts_table(tmp_path, monkeypatch) -> None:
    from openpyxl import Workbook

    target = _kb_dir(tmp_path) / "pnl.xlsx"
    wb = Workbook()
    ws = wb.active
    ws["A1"] = "REVENUE_MARKER"
    ws["B1"] = 12345
    wb.save(str(target))

    client = _client(tmp_path, monkeypatch)
    res = _preview(client, "knowledge/Downloads/pnl.xlsx")

    assert res.status_code == 200, res.text
    body = res.json()
    assert "REVENUE_MARKER" in body["content"]
    assert "12345" in body["content"]
    assert body["extracted"] is True
    assert body["sourceTool"] == "xlsx_read"


def test_korean_named_docx_preview_extracts(tmp_path, monkeypatch) -> None:
    # The real report that surfaced the bug had a Korean filename; confirm the
    # cause was the binary read, not filename encoding — a Korean-named docx
    # previews fine once extraction is wired.
    from docx import Document

    name = "내외디스틸러리 (주) 농업회사법인 (3).docx"
    target = _kb_dir(tmp_path) / name
    doc = Document()
    doc.add_paragraph("한글 재무 요약 UNIQUE_KO")
    doc.save(str(target))

    client = _client(tmp_path, monkeypatch)
    res = _preview(client, f"knowledge/Downloads/{name}")

    assert res.status_code == 200, res.text
    assert "UNIQUE_KO" in res.json()["content"]


# --------------------------------------------------------------------------- #
# Text-like files keep the fast path (regression)
# --------------------------------------------------------------------------- #
def test_text_file_preview_unchanged(tmp_path, monkeypatch) -> None:
    target = _kb_dir(tmp_path) / "notes.md"
    target.write_text("# Heading\n\nplain markdown body", encoding="utf-8")

    client = _client(tmp_path, monkeypatch)
    res = _preview(client, "knowledge/Downloads/notes.md")

    assert res.status_code == 200, res.text
    body = res.json()
    assert body["content"] == "# Heading\n\nplain markdown body"
    # Plain read path does not set the extraction flags.
    assert "extracted" not in body


# --------------------------------------------------------------------------- #
# Honest errors — never a misleading not_found for a present-but-unreadable file
# --------------------------------------------------------------------------- #
def test_corrupt_pdf_yields_honest_error_not_not_found(tmp_path, monkeypatch) -> None:
    target = _kb_dir(tmp_path) / "broken.pdf"
    # Genuinely non-UTF-8 bytes (like a real PDF) so the fast text read fails and
    # the extraction path runs — then pypdf rejects the malformed document.
    target.write_bytes(b"%PDF-1.4\n\xff\xfe\x80 binary \xc0\xc1 garbage")

    client = _client(tmp_path, monkeypatch)
    res = _preview(client, "knowledge/Downloads/broken.pdf")

    assert res.status_code == 415, res.text
    body = res.json()
    assert body["error"] != "not_found"
    assert "preview unavailable" in body["error"]


def test_missing_file_still_not_found(tmp_path, monkeypatch) -> None:
    _kb_dir(tmp_path)  # collection exists, file does not
    client = _client(tmp_path, monkeypatch)
    res = _preview(client, "knowledge/Downloads/ghost.pdf")

    assert res.status_code == 404
    assert res.json()["error"] == "not_found"


# --------------------------------------------------------------------------- #
# Companion fixes
# --------------------------------------------------------------------------- #
def test_kb_context_convertible_includes_legacy_xls() -> None:
    # The chat-attachment inlining path shares the same converter; legacy .xls
    # must inline, not fall through to the unsupported-format note.
    from magi_agent.transport.kb_context import _CONVERTIBLE_EXTENSIONS

    assert ".xls" in _CONVERTIBLE_EXTENSIONS
    assert ".xlsx" in _CONVERTIBLE_EXTENSIONS


def test_document_extraction_libs_are_base_dependencies() -> None:
    # The preview is a first-party local reading surface: its extraction
    # libraries must ship by default (not gated behind the optional [files]
    # extra), so `pip install magi-agent` previews binary docs out-of-the-box.
    pyproject = Path(__file__).resolve().parents[1] / "pyproject.toml"
    data = tomllib.loads(pyproject.read_text(encoding="utf-8"))
    core = data["project"]["dependencies"]
    core_names = {_dist_name(spec) for spec in core}
    for name in ("pypdf", "openpyxl", "xlrd", "python-pptx"):
        assert name in core_names, f"{name} must be a base dependency"
    # The old [files] extra is now an empty back-compat stub.
    assert data["project"]["optional-dependencies"]["files"] == []


def _dist_name(spec: str) -> str:
    out = []
    for ch in spec.strip():
        if ch.isalnum() or ch in "-_.":
            out.append(ch)
        else:
            break
    return "".join(out)

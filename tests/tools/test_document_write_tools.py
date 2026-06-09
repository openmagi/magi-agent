"""Tests for the DOCX DocumentWrite backend and the format dispatch.

Hermetic: the ``files`` extra installs ``python-docx`` so the happy-path test
can create and re-read a real ``.docx``.  The dependency-not-installed path is
simulated with ``patch.dict("sys.modules", {"docx": None})``.
"""

from __future__ import annotations

import subprocess
import sys
from pathlib import Path
from unittest.mock import patch

from magi_agent.plugins.native.documents import document_write
from magi_agent.tools.context import ToolContext
from magi_agent.tools.document_write_tools import docx_write

_MARKDOWN = """# Quarterly Report

This is the opening paragraph with **bold** and *italic* emphasis.

## Highlights

- First bullet item
- Second bullet item

## Steps

1. First ordered step
2. Second ordered step

## Metrics

| Metric | Value |
| --- | --- |
| Revenue | 1000 |
| Growth | 12 |

```
print("hello world")
```
"""


def _context(workspace_root: Path) -> ToolContext:
    return ToolContext(
        botId="bot-doc-write-test",
        sessionId="session-doc-write-test",
        turnId="turn-doc-write-test",
        workspaceRoot=str(workspace_root),
    )


def _docx_text(path: Path) -> str:
    from docx import Document  # type: ignore[import]

    doc = Document(str(path))
    parts: list[str] = [p.text for p in doc.paragraphs]
    for table in doc.tables:
        for row in table.rows:
            parts.extend(cell.text for cell in row.cells)
    return "\n".join(parts)


class TestDocxWriteHappyPath:
    def test_writes_valid_docx_with_all_constructs(self, tmp_path: Path) -> None:
        result = docx_write(
            {"content": _MARKDOWN, "path": "report.docx"}, _context(tmp_path)
        )

        assert result.status == "ok", result.error_code
        out = result.output
        assert out["format"] == "docx"
        assert out["path"] == "report.docx"
        assert out["pathRef"] == "report.docx"
        assert isinstance(out["contentDigest"], str) and out["contentDigest"]
        assert isinstance(out["byteCount"], int) and out["byteCount"] > 0
        assert out["artifactRefs"][0].startswith("artifact:docx:")
        assert isinstance(out["artifactRef"], str) and out["artifactRef"].startswith("artifact:docx:")
        assert out["artifactRef"] == out["artifactRefs"][0]
        assert result.artifact_refs == tuple(out["artifactRefs"])

        saved = tmp_path / "report.docx"
        assert saved.exists()

        text = _docx_text(saved)
        # Heading text present.
        assert "Quarterly Report" in text
        assert "Highlights" in text
        # Paragraph (inline emphasis text preserved, markers stripped).
        assert "opening paragraph" in text
        assert "bold" in text
        assert "italic" in text
        # Bullet + ordered items.
        assert "First bullet item" in text
        assert "Second ordered step" in text
        # Table cell content.
        assert "Revenue" in text
        assert "1000" in text
        # Code block content preserved.
        assert 'print("hello world")' in text

    def test_default_filename(self, tmp_path: Path) -> None:
        result = docx_write({"content": "# Title\n\nBody."}, _context(tmp_path))
        assert result.status == "ok"
        assert (tmp_path / "magi-document.docx").exists()

    def test_empty_content_blocked(self, tmp_path: Path) -> None:
        result = docx_write({"content": "   "}, _context(tmp_path))
        assert result.status == "blocked"
        assert result.error_code == "content_required"


class TestDocxWriteDependency:
    def test_missing_docx_returns_blocked(self, tmp_path: Path) -> None:
        with patch.dict("sys.modules", {"docx": None}):
            result = docx_write(
                {"content": "# X\n\nbody", "path": "a.docx"}, _context(tmp_path)
            )
        assert result.status == "blocked"
        assert result.error_code == "document_dependency_not_installed"


class TestDocxWritePathEscape:
    def test_path_traversal_blocked(self, tmp_path: Path) -> None:
        result = docx_write(
            {"content": "# X\n\nbody", "path": "../escape.docx"}, _context(tmp_path)
        )
        assert result.status == "blocked"
        assert result.error_code in {
            "path_traversal_blocked",
            "absolute_path_blocked",
            "hidden_path_write_blocked",
        }

    def test_absolute_path_blocked(self, tmp_path: Path) -> None:
        result = docx_write(
            {"content": "# X\n\nbody", "path": "/etc/evil.docx"}, _context(tmp_path)
        )
        assert result.status == "blocked"
        assert result.error_code == "absolute_path_blocked"


class TestDocumentWriteDispatch:
    def test_format_docx_routes_to_docx(self, tmp_path: Path) -> None:
        result = document_write(
            {"content": _MARKDOWN, "path": "dispatch.docx", "format": "docx"},
            _context(tmp_path),
        )
        assert result.status == "ok"
        assert result.output["format"] == "docx"
        saved = tmp_path / "dispatch.docx"
        assert saved.exists()
        assert "Quarterly Report" in _docx_text(saved)

    def test_docx_suffix_routes_to_docx_without_format_arg(
        self, tmp_path: Path
    ) -> None:
        result = document_write(
            {"content": "# Heading\n\nbody text", "path": "by-suffix.docx"},
            _context(tmp_path),
        )
        assert result.status == "ok"
        assert result.output["format"] == "docx"
        assert (tmp_path / "by-suffix.docx").exists()

    def test_md_path_unchanged_raw_markdown(self, tmp_path: Path) -> None:
        result = document_write(
            {"content": "# Heading\n\nraw body", "path": "note.md"},
            _context(tmp_path),
        )
        assert result.status == "ok"
        # Existing markdown path: no ``format`` key, raw text written verbatim.
        assert "format" not in result.output
        saved = tmp_path / "note.md"
        assert saved.exists()
        assert saved.read_text(encoding="utf-8") == "# Heading\n\nraw body"


class TestDocxWriteRedactionContract:
    """M6: pin that the DOCX output contains the redacted form, not the raw token."""

    def test_private_path_is_redacted_in_output(self, tmp_path: Path) -> None:
        raw_path = "/home/user/secret.py"
        source = f"# Report\n\nSee {raw_path} for details."
        result = docx_write(
            {"content": source, "path": "redact-test.docx"}, _context(tmp_path)
        )

        assert result.status == "ok", result.error_code
        saved = tmp_path / "redact-test.docx"
        text = _docx_text(saved)

        # (a) Redacted form is present in the output.
        assert "[redacted-path]" in text
        # (b) Raw token is NOT present — Task B must compare against redacted source.
        assert raw_path not in text


class TestDocxWriteTruncation:
    """M7: a source exceeding max-chars is truncated without raising and returns ok."""

    def test_oversized_source_truncated_without_error(self, tmp_path: Path) -> None:
        # Build a source that exceeds the 200_000-char cap.
        long_source = "# Title\n\n" + ("word " * 50_000)
        assert len(long_source) > 200_000

        result = docx_write(
            {"content": long_source, "path": "long-doc.docx"}, _context(tmp_path)
        )

        assert result.status == "ok", result.error_code
        assert (tmp_path / "long-doc.docx").exists()
        assert result.output["byteCount"] > 0


class TestDocxWriteCoverageEvidence:
    """Task B: docx_write emits a digest-only DocumentCoverage evidence record."""

    def test_happy_path_emits_pass_coverage_via_collector(self, tmp_path: Path) -> None:
        from magi_agent.evidence.local_tool_collector import LocalToolEvidenceCollector

        result = docx_write(
            {"content": _MARKDOWN, "path": "cov.docx"}, _context(tmp_path)
        )
        assert result.status == "ok", result.error_code

        # Coverage projection is visible in the tool output.
        coverage = result.output["coverage"]
        assert coverage["type"] == "DocumentCoverage"
        assert coverage["status"] == "pass"
        assert coverage["coverageRatio"] >= 0.95
        # Digest-only: only digests/numbers, no raw words.
        assert coverage["sourceDigest"].startswith("sha256:")
        assert coverage["docDigest"].startswith("sha256:")
        assert "Quarterly Report" not in repr(coverage)

        # The established emission channel: metadata["evidence"] → collector →
        # canonical EvidenceRecord consumed by the verifier-bus (Task C).
        assert "evidence" in result.metadata
        collector = LocalToolEvidenceCollector()
        records = collector.record_tool_result(
            session_id="s1",
            turn_id="t1",
            tool_call_id="call-docx",
            tool_name="DocumentWrite",
            result=result,
        )
        coverage_records = [
            r for r in records if getattr(r, "type", None) == "DocumentCoverage"
        ]
        assert len(coverage_records) == 1
        evidence = coverage_records[0]
        assert evidence.status == "ok"
        assert evidence.fields["status"] == "pass"

    def test_mismatched_render_emits_failed_coverage_but_tool_ok(
        self, tmp_path: Path
    ) -> None:
        # A pipe-heavy source that is NOT a valid table (no separator row) renders
        # each line as a plain paragraph, but the boundary still tokenizes the
        # words; construct a case where rendering drops content by feeding source
        # that the renderer collapses. Easiest deterministic mismatch: monkeypatch
        # the coverage to compare a faithful source against truncated doc text.
        from unittest.mock import patch

        from magi_agent.tools import document_tools

        original = document_tools.extract_docx_text

        def _truncating_extract(doc: object) -> str:
            # Drop everything after the first paragraph to force missing units.
            full = original(doc)
            return full.split("\n\n")[0]

        with patch.object(document_tools, "extract_docx_text", _truncating_extract):
            result = docx_write(
                {"content": _MARKDOWN, "path": "mismatch.docx"}, _context(tmp_path)
            )

        # Tool still succeeds (audit-only — no blocking in Task B).
        assert result.status == "ok", result.error_code
        coverage = result.output["coverage"]
        assert coverage["status"] == "failed"
        assert coverage["coverageRatio"] < 1.0
        assert len(coverage["missingUnitDigests"]) > 0
        # Evidence record reflects the failed coverage.
        assert result.metadata["evidence"]["status"] == "failed"

    def test_redaction_contract_coverage_passes_against_redacted_source(
        self, tmp_path: Path
    ) -> None:
        # Source contains a private path that is redacted before rendering.
        # Coverage compares against the redacted safe_source, so the redacted
        # token (present in the rendered doc) counts as covered → pass.
        source = "# Report\n\nSee /home/user/secret.py for details."
        result = docx_write(
            {"content": source, "path": "redact-cov.docx"}, _context(tmp_path)
        )
        assert result.status == "ok"
        coverage = result.output["coverage"]
        assert coverage["status"] == "pass"
        assert coverage["coverageRatio"] == 1.0


class TestDocxWriteDualStatusContract:
    """Fix 2: pin the dual-status contract explicitly.

    When coverage is "failed" but the write succeeds, ``EvidenceRecord.status``
    (top-level) must be ``"ok"`` (following ``ToolResult.status``), while the
    coverage verdict lives in ``EvidenceRecord.fields["status"] == "failed"``.
    Consumers that need to block on failed coverage MUST check ``fields["status"]``.
    """

    def test_failed_coverage_via_collector_ok_top_level_failed_fields(
        self, tmp_path: Path
    ) -> None:
        from unittest.mock import patch

        from magi_agent.evidence.local_tool_collector import LocalToolEvidenceCollector
        from magi_agent.tools import document_tools

        original = document_tools.extract_docx_text

        def _truncating_extract(doc: object) -> str:
            # Return only the first paragraph to force many missing units.
            full = original(doc)
            return full.split("\n\n")[0] if "\n\n" in full else full[:20]

        with patch.object(document_tools, "extract_docx_text", _truncating_extract):
            result = docx_write(
                {"content": _MARKDOWN, "path": "dual-status.docx"}, _context(tmp_path)
            )

        # Tool write succeeded.
        assert result.status == "ok", result.error_code
        # Coverage projection shows failed.
        assert result.output["coverage"]["status"] == "failed"
        # Collect through the production path.
        assert "evidence" in result.metadata
        collector = LocalToolEvidenceCollector()
        records = collector.record_tool_result(
            session_id="s1",
            turn_id="t1",
            tool_call_id="call-dual",
            tool_name="DocumentWrite",
            result=result,
        )
        coverage_records = [
            r for r in records if getattr(r, "type", None) == "DocumentCoverage"
        ]
        assert len(coverage_records) == 1
        evidence = coverage_records[0]
        # Dual-status contract: top-level status follows ToolResult.status ("ok").
        assert evidence.status == "ok", (
            f"Expected evidence.status=='ok' (follows ToolResult.status), got {evidence.status!r}. "
            "Consumers checking coverage verdict MUST use fields['status']."
        )
        # Coverage verdict lives in fields["status"].
        assert evidence.fields["status"] == "failed", (
            f"Expected fields['status']=='failed', got {evidence.fields['status']!r}"
        )


class TestDocumentWriteImportBoundary:
    def test_document_write_tools_import_does_not_load_docx(self) -> None:
        completed = subprocess.run(
            [
                sys.executable,
                "-c",
                (
                    "import importlib, sys; "
                    "importlib.import_module('magi_agent.tools.document_write_tools'); "
                    "assert 'docx' not in sys.modules, "
                    "'document_write_tools import pulled docx into sys.modules'"
                ),
            ],
            capture_output=True,
            text=True,
            check=False,
        )
        assert completed.returncode == 0, completed.stderr

    def test_documents_plugin_import_does_not_load_docx(self) -> None:
        completed = subprocess.run(
            [
                sys.executable,
                "-c",
                (
                    "import importlib, sys; "
                    "importlib.import_module('magi_agent.plugins.native.documents'); "
                    "assert 'docx' not in sys.modules, "
                    "'documents plugin import pulled docx into sys.modules'"
                ),
            ],
            capture_output=True,
            text=True,
            check=False,
        )
        assert completed.returncode == 0, completed.stderr
